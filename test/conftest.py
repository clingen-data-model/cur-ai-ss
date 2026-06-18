import io
import os
import random
from pathlib import Path
from typing import Optional

import openpyxl
import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image
from sqlalchemy.orm import Session

from lib.api import db
from lib.core.environment import env
from lib.models import AgentRunDB, Base


@pytest.fixture
def test_resources_path():
    return os.path.abspath(os.path.join(os.path.dirname(__file__), 'resources'))


@pytest.fixture
def test_file_contents(test_resources_path):
    def _loader(file_name, mode='r'):
        with open(os.path.join(test_resources_path, file_name), mode) as file:
            return file.read()

    return _loader


@pytest.fixture
def mocked_root_dir(monkeypatch, tmpdir):
    """Mock CAA_ROOT to point to a temporary directory."""
    monkeypatch.setattr(db.env, 'CAA_ROOT', str(tmpdir))
    return tmpdir


@pytest.fixture
def db_session(mocked_root_dir, monkeypatch):
    db.env.SQLLITE_DIR = ''
    monkeypatch.setattr(db, '_engine', None)
    monkeypatch.setattr(db, '_session_factory', None)
    engine = db.get_engine()
    Base.metadata.create_all(bind=engine)
    session_local = db.get_sessionmaker()
    session: Session = session_local()
    yield session
    session.rollback()
    session.close()


@pytest.fixture
def agent_run(db_session):
    run = AgentRunDB(
        git_hash='abc123def456',
        description='test run',
        model=env.OPENAI_API_DEPLOYMENT,
    )
    db_session.add(run)
    db_session.flush()
    return run


@pytest.fixture
def docx_with_image():
    doc = Document()
    doc.add_heading('Test Document with Image', 0)
    doc.add_paragraph('This is a test DOCX document with an embedded image.')

    # Create a simple test image
    img = Image.new('RGB', (100, 100), color=(73, 109, 137))
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)

    doc.add_picture(img_bytes, width=Inches(2))
    doc.add_paragraph('Image caption text.')

    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


@pytest.fixture
def xlsx_with_data():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Sheet1'
    ws['A1'] = 'Name'
    ws['B1'] = 'Value'
    ws['A2'] = 'Alpha'
    ws['B2'] = 42
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()
